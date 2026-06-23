#!/usr/bin/env python3
"""
生成 pandoc reference.docx 样式模板
修复：标题居中、页眉页脚、表格、右对齐落款
"""
import subprocess, os, sys, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HR_DIR = os.path.dirname(os.path.abspath(__file__))
REF_DOC = os.path.join(HR_DIR, "reference.docx")
COMPANY = "广州进格智能科技有限公司"

# ── 工具函数 ──────────────────────────────────────────

def black(style):
    """把样式字体颜色强制设为黑色（覆盖 pandoc 默认蓝色标题）"""
    try:
        style.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass

def add_style(doc, name, base="Normal", alignment=None, size=None, bold=None):
    """新增（或获取已有）段落样式"""
    from docx.enum.style import WD_STYLE_TYPE
    try:
        s = doc.styles[name]
    except KeyError:
        s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        try:
            s.base_style = doc.styles[base]
        except Exception:
            pass
    if alignment is not None:
        s.paragraph_format.alignment = alignment
    if size is not None:
        s.font.size = Pt(size)
    if bold is not None:
        s.font.bold = bold
    return s

def set_page_header(doc, text):
    """在所有节中设置页眉文字，居中，不追加装饰边框"""
    for section in doc.sections:
        section.different_first_page_header_footer = False
        header = section.header
        for p in header.paragraphs:
            p.clear()
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()
        run = para.add_run(text)
        run.font.size = Pt(9)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ppr = para._p.get_or_add_pPr()
        p_bdr = ppr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            ppr.append(p_bdr)
        bottom = p_bdr.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            p_bdr.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")

def set_page_footer(doc):
    """在所有节中设置居中页码页脚"""
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            p.clear()
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("第 ")
        run.font.size = Pt(9)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)

        run2 = para.add_run(" 页")
        run2.font.size = Pt(9)

def resolve_pandoc():
    candidates = [
        os.path.expanduser("~/bin/pandoc"),
        "/opt/homebrew/bin/pandoc",
        "/usr/local/bin/pandoc",
        shutil.which("pandoc"),
    ]
    for candidate in candidates:
        if candidate and os.path.isfile(candidate) and os.access(candidate, os.X_OK):
            return candidate
    return None

def fix_table_borders(doc):
    """给 reference.docx 里所有表格加上清晰边框"""
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top","left","bottom","right","insideH","insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            tblBorders.append(el)
        # 移除旧 borders 再插入
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(tblBorders)

# ── 主流程 ────────────────────────────────────────────

def main():
    pandoc = resolve_pandoc()
    if not pandoc:
        print("❌ 未找到 pandoc，请检查安装路径")
        sys.exit(1)

    # 1. 生成 pandoc 默认 reference.docx
    result = subprocess.run(
        [pandoc, "--print-default-data-file", "reference.docx"],
        capture_output=True
    )
    if result.returncode != 0 or not result.stdout:
        print("❌ 无法生成默认 reference.docx，请检查 pandoc 路径")
        sys.exit(1)
    with open(REF_DOC, "wb") as f:
        f.write(result.stdout)

    doc = Document(REF_DOC)

    # 2. 修改 Heading 1 → 居中、黑色
    for style in doc.styles:
        if style.name == "Heading 1":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.font.size = Pt(16)
            style.font.bold = True
            black(style)
        elif style.name in ("Heading 2", "Heading 3"):
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style.font.size = Pt(12 if style.name == "Heading 2" else 11)
            style.font.bold = True
            black(style)

    # 3. 新增 Para-Right（右对齐，用于落款）
    add_style(doc, "Para-Right", base="Normal",
              alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # 4. 新增 Para-Center（居中，用于需要居中的正文）
    add_style(doc, "Para-Center", base="Normal",
              alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 5. 新增 CoverTitle（封面大标题：28pt，居中，加粗）
    add_style(doc, "CoverTitle", base="Normal",
              alignment=WD_ALIGN_PARAGRAPH.CENTER, size=28, bold=True)

    # 5. 修复表格边框
    fix_table_borders(doc)

    # 6. 设置页眉页脚
    set_page_header(doc, COMPANY)
    set_page_footer(doc)

    doc.save(REF_DOC)
    print(f"✅ reference.docx 已生成：{REF_DOC}")

if __name__ == "__main__":
    main()
